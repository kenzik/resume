#!/usr/bin/env python3

import argparse
import sys
import os
import yaml
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
from fpdf.enums import XPos, YPos


# --- CONSTANTS ---
# Section titles that should not be displayed as headers
SKIP_SECTION_TITLES = {"Highlights", "Key Projects"}

# Resume section titles (used across all formats)
SECTION_PROFILE = "PROFILE"
SECTION_COMPETENCIES = "CORE COMPETENCIES"
SECTION_EXPERIENCE = "PROFESSIONAL EXPERIENCE"
SECTION_EARLIER = "EARLIER ROLES"
SECTION_CERTS_EDU = "CERTIFICATIONS & EDUCATION"

# Separators
SEPARATOR_CONTACT = " | "  # Contact info separator
SEPARATOR_SKILLS_DOCX = " • "  # Skills separator for DOCX/MD
SEPARATOR_SKILLS_PDF = " - "  # Skills separator for PDF
SEPARATOR_JOB = " | "  # Job company | title separator
SEPARATOR_MD_SECTION = "---"  # Markdown section separator

# DOCX constants
DOCX_FONT_NAME = "Calibri"
DOCX_FONT_SIZE = 11

# PDF layout constants
PDF_FONT_FAMILY = "Helvetica"
PDF_FOOTER_MARGIN = -15
PDF_FOOTER_FONT_SIZE = 8
PDF_PAGE_BREAK_Y = 250
PDF_MARGIN = 15
PDF_HEADER_FONT_SIZE = 24
PDF_SECTION_FONT_SIZE = 12
PDF_BODY_FONT_SIZE = 10
PDF_JOB_TITLE_FONT_SIZE = 11
PDF_SUBSECTION_FONT_SIZE = 10
PDF_TECH_FONT_SIZE = 9
PDF_BULLET_PREFIX = "- "  # Bullet point prefix for PDF

# RTF constants
RTF_BULLET = "\u2022"  # Proper bullet character
RTF_FONT_NAME = "Arial"  # RTF font name
RTF_HEADER = r"{\rtf1\ansi\deff0{\fonttbl{\f0 Arial;}}\viewkind4\uc1\pard\sa200\sl276\slmult1\lang9\f0\fs22"

# Education delimiter
EDU_DELIMITER = "—"

# Magic strings for formatting logic
STRING_IN_PROGRESS = "In Progress"
STRING_UNIVERSITY = "University"
STRING_REDACTED_LOC = "[REDACTED LOC]"
STRING_REDACTED_TEL = "[REDACTED TEL]"


# --- UTILS ---
# Pre-compile translation table for PDF cleaning (much faster than multiple replace calls)
_PDF_REPLACEMENTS = str.maketrans(
    {
        "\u2022": "-",
        "\u2013": "-",
        "\u2014": "--",
        "\u2018": "'",
        "\u2019": "'",
        "\u201c": '"',
        "\u201d": '"',
        "\u2026": "...",
    }
)


def clean_pdf(text):
    """Clean text for PDF output using efficient str.translate()."""
    if not text:
        return ""
    text = text.translate(_PDF_REPLACEMENTS)
    return text.encode("latin-1", "replace").decode("latin-1")


def clean_rtf(text):
    """Clean text for RTF output with proper escaping."""
    if not text:
        return ""
    # Escape special RTF characters first
    text = text.replace("\\", "\\\\").replace("{", "\\{").replace("}", "\\}")
    # Process non-ASCII characters more efficiently
    res = []
    for char in text:
        code = ord(char)
        if code < 128:
            res.append(char)
        else:
            # RTF Unicode escape: \uN? where N is signed 16-bit integer
            rtf_code = code - 65536 if code > 32767 else code
            res.append(f"\\u{rtf_code}?")
    return "".join(res)


def get_contact_string(reveal_pii):
    """
    Constructs the contact string from Env Vars.
    Default: Redacts City/Phone.
    --reveal-pii: Shows everything.
    """
    city = os.getenv("RESUME_CITY_STATE", "City, State")
    phone = os.getenv("RESUME_PHONE", "555-555-5555")
    email = os.getenv("RESUME_EMAIL", "email@example.com")
    linkedin = os.getenv("RESUME_LINKEDIN", "linkedin.com/in/user")

    parts = []

    if reveal_pii:
        parts.append(city)
        parts.append(phone)
    else:
        parts.append(STRING_REDACTED_LOC)
        parts.append(STRING_REDACTED_TEL)

    parts.append(email)
    parts.append(linkedin)

    return SEPARATOR_CONTACT.join(parts)


# --- HELPER FUNCTIONS ---
def should_show_section_title(title):
    """Check if a section title should be displayed as a header."""
    return title and title not in SKIP_SECTION_TITLES


def parse_education(edu_string):
    """Parse education string into institution and degree parts."""
    parts = edu_string.split(EDU_DELIMITER, 1)
    institution = parts[0].strip()
    degree = parts[1].strip() if len(parts) > 1 else None
    return institution, degree


# --- RENDERERS ---


def render_docx(data, filename):
    print(f"[+] Generating DOCX -> {filename}...")
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = DOCX_FONT_NAME
    style.font.size = Pt(DOCX_FONT_SIZE)

    # Header
    h1 = doc.add_heading(data["header"]["name"], 0)
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.add_run(data["header"]["contact"]).bold = True

    # Sections
    doc.add_heading(SECTION_PROFILE, level=1)
    doc.add_paragraph(data["profile"])

    doc.add_heading(SECTION_COMPETENCIES, level=1)
    doc.add_paragraph(SEPARATOR_SKILLS_DOCX.join(data["skills"]))

    doc.add_heading(SECTION_EXPERIENCE, level=1)
    for job in data["experience"]:
        p = doc.add_paragraph()
        p.add_run(job["company"]).bold = True
        p.add_run(f"{SEPARATOR_JOB}{job['title']}").italic = True
        p.add_run(f"\n{job['date']}")

        if job.get("summary"):
            doc.add_paragraph(job["summary"])

        for section in job.get("sections", []):
            if should_show_section_title(section.get("title")):
                doc.add_paragraph(section["title"]).bold = True
            for bullet in section["bullets"]:
                doc.add_paragraph(bullet, style="List Bullet")

        if job.get("tech"):
            doc.add_paragraph(job["tech"]).italic = True
        doc.add_paragraph()

    if data.get("earlier"):
        doc.add_heading(SECTION_EARLIER, level=1)
        for role in data["earlier"]:
            doc.add_paragraph(role)

    doc.add_heading(SECTION_CERTS_EDU, level=1)
    for cert in data["certs"]:
        p = doc.add_paragraph()
        p.add_run(cert).bold = STRING_IN_PROGRESS in cert
    for edu in data["education"]:
        institution, degree = parse_education(edu)
        p = doc.add_paragraph()
        p.add_run(f"{institution} {EDU_DELIMITER} ").bold = True
        if degree:
            p.add_run(degree)

    doc.save(filename)


def render_md(data, filename):
    """Render resume as Markdown with optimized string building."""
    print(f"[+] Generating MD -> {filename}...")
    try:
        lines = [
            f"# {data['header']['name']}\n",
            f"**{data['header']['contact']}**\n",
            f"## {SECTION_PROFILE}\n\n{data['profile']}\n",
            f"## {SECTION_COMPETENCIES}\n\n{SEPARATOR_SKILLS_DOCX.join(data['skills'])}\n",
            f"## {SECTION_EXPERIENCE}\n",
        ]

        for job in data["experience"]:
            lines.append(
                f"### {job['company']}{SEPARATOR_JOB}*{job['title']}*\n**{job['date']}**\n"
            )
            if job.get("summary"):
                lines.append(f"{job['summary']}\n")
            for section in job.get("sections", []):
                if should_show_section_title(section.get("title")):
                    lines.append(f"**{section['title']}**\n")
                lines.extend(f"* {bullet}\n" for bullet in section["bullets"])
                lines.append("\n")
            if job.get("tech"):
                lines.append(f"*{job['tech']}*\n")
            lines.append(f"{SEPARATOR_MD_SECTION}\n")

        if data.get("earlier"):
            lines.append(f"\n## {SECTION_EARLIER}\n")
            lines.extend(f"* {role}\n" for role in data["earlier"])
        lines.append(f"\n## {SECTION_CERTS_EDU}\n")
        lines.extend(f"* **{cert}**\n" for cert in data["certs"])
        for edu in data["education"]:
            institution, degree = parse_education(edu)
            if degree:
                lines.append(f"* **{institution}** {EDU_DELIMITER} {degree}\n")
            else:
                lines.append(f"* **{institution}**\n")

        with open(filename, "w", encoding="utf-8") as f:
            f.writelines(lines)
    except Exception as e:
        print(f"FAILED to write Markdown: {e}")
        raise e


def render_pdf(data, filename):
    print(f"[+] Generating PDF -> {filename}...")

    class PDF(FPDF):
        def footer(self):
            self.set_y(PDF_FOOTER_MARGIN)
            self.set_font(PDF_FONT_FAMILY, "I", PDF_FOOTER_FONT_SIZE)
            self.cell(0, 10, f"Page {self.page_no()}", align="C")

    pdf = PDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=PDF_MARGIN)
    width = pdf.epw

    # --- HEADER ---
    pdf.set_font(PDF_FONT_FAMILY, "B", PDF_HEADER_FONT_SIZE)
    pdf.cell(
        width,
        10,
        clean_pdf(data["header"]["name"]),
        new_x=XPos.LMARGIN,
        new_y=YPos.NEXT,
        align="C",
    )
    pdf.set_font(PDF_FONT_FAMILY, "B", PDF_BODY_FONT_SIZE)  # Bold contact info
    pdf.cell(
        width,
        10,
        clean_pdf(data["header"]["contact"]),
        new_x=XPos.LMARGIN,
        new_y=YPos.NEXT,
        align="C",
    )
    pdf.ln(5)

    # --- SECTIONS ---
    sections = [
        (SECTION_PROFILE, "text", data["profile"]),
        (SECTION_COMPETENCIES, "text", SEPARATOR_SKILLS_PDF.join(data["skills"])),
    ]

    for title, mode, content in sections:
        pdf.set_font(PDF_FONT_FAMILY, "B", PDF_SECTION_FONT_SIZE)
        pdf.set_fill_color(220, 220, 220)
        pdf.cell(width, 8, title, new_x=XPos.LMARGIN, new_y=YPos.NEXT, fill=False)
        pdf.ln(2)
        pdf.set_font(PDF_FONT_FAMILY, "", PDF_BODY_FONT_SIZE)
        pdf.multi_cell(width, 5, clean_pdf(content))
        pdf.ln(4)

    # --- EXPERIENCE ---
    pdf.set_font(PDF_FONT_FAMILY, "B", PDF_SECTION_FONT_SIZE)
    pdf.cell(width, 8, SECTION_EXPERIENCE, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(2)

    for job in data["experience"]:
        pdf.set_font(PDF_FONT_FAMILY, "B", PDF_JOB_TITLE_FONT_SIZE)
        pdf.cell(
            width,
            6,
            clean_pdf(f"{job['company']}{SEPARATOR_JOB}{job['title']}"),
            new_x=XPos.LMARGIN,
            new_y=YPos.NEXT,
        )
        pdf.set_font(PDF_FONT_FAMILY, "", PDF_BODY_FONT_SIZE)
        pdf.cell(width, 5, clean_pdf(job["date"]), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(2)

        if job.get("summary"):
            pdf.multi_cell(width, 5, clean_pdf(job["summary"]))
            pdf.ln(2)

        for section in job.get("sections", []):
            if should_show_section_title(section.get("title")):
                pdf.set_font(PDF_FONT_FAMILY, "B", PDF_SUBSECTION_FONT_SIZE)
                pdf.cell(
                    width,
                    6,
                    clean_pdf(section["title"]),
                    new_x=XPos.LMARGIN,
                    new_y=YPos.NEXT,
                )
            pdf.set_font(PDF_FONT_FAMILY, "", PDF_BODY_FONT_SIZE)
            for bullet in section["bullets"]:
                pdf.set_x(pdf.l_margin + 2)
                pdf.multi_cell(width - 2, 5, f"{PDF_BULLET_PREFIX}{clean_pdf(bullet)}")
            pdf.ln(2)

        if job.get("tech"):
            pdf.set_font(PDF_FONT_FAMILY, "I", PDF_TECH_FONT_SIZE)
            pdf.multi_cell(width, 5, clean_pdf(job["tech"]))
        pdf.ln(4)

    # --- TAIL SECTIONS ---
    if data.get("earlier"):
        if pdf.get_y() > PDF_PAGE_BREAK_Y:
            pdf.add_page()
        pdf.set_font(PDF_FONT_FAMILY, "B", PDF_SECTION_FONT_SIZE)
        pdf.cell(width, 8, SECTION_EARLIER, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(2)
        pdf.set_font(PDF_FONT_FAMILY, "", PDF_BODY_FONT_SIZE)
        for role in data["earlier"]:
            pdf.multi_cell(width, 5, clean_pdf(role), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(4)

    if pdf.get_y() > PDF_PAGE_BREAK_Y:
        pdf.add_page()
    pdf.set_font(PDF_FONT_FAMILY, "B", PDF_SECTION_FONT_SIZE)
    pdf.cell(width, 8, SECTION_CERTS_EDU, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(2)
    for item in data["certs"] + data["education"]:
        is_bold = STRING_IN_PROGRESS in item or STRING_UNIVERSITY in item
        pdf.set_font(PDF_FONT_FAMILY, "B" if is_bold else "", PDF_BODY_FONT_SIZE)
        pdf.multi_cell(width, 6, clean_pdf(item), new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf.output(filename)


def render_rtf(data, filename):
    """Render resume as RTF with proper Unicode handling."""
    print(f"[+] Generating RTF -> {filename}...")
    rtf = [RTF_HEADER]

    def line(txt, bold=False, italic=False, ul=False):
        fmt = r"\par"
        if bold:
            fmt = r"\b " + fmt + r"\b0"
        if italic:
            fmt = r"\i " + fmt + r"\i0"
        if ul:
            fmt = r"\ul " + fmt + r"\ulnone"
        rtf.append(
            r"\pard\sa200\sl276\slmult1 "
            + fmt.replace(r"\par", clean_rtf(txt) + r"\par")
        )

    # Header
    rtf.append(r"\qc\b\fs32 " + clean_rtf(data["header"]["name"]) + r"\par")
    rtf.append(r"\fs22\b " + clean_rtf(data["header"]["contact"]) + r"\b0\par\par")

    line(SECTION_PROFILE, bold=True, ul=True)
    line(data["profile"])
    line(SECTION_COMPETENCIES, bold=True, ul=True)
    line(f" {RTF_BULLET} ".join(data["skills"]))

    line(SECTION_EXPERIENCE, bold=True, ul=True)
    for job in data["experience"]:
        rtf.append(
            r"\pard\sa200\sl276\slmult1\b "
            + clean_rtf(job["company"])
            + r"\b0 "
            + clean_rtf(SEPARATOR_JOB)
            + r"\i "
            + clean_rtf(job["title"])
            + r"\i0\par"
        )
        line(job["date"])
        if job.get("summary"):
            line(job["summary"])
        for section in job.get("sections", []):
            if should_show_section_title(section.get("title")):
                line(section["title"], bold=True)
            for bullet in section["bullets"]:
                rtf.append(
                    r"\pard\sa200\sl276\slmult1\tx360\li360\fi-360 "
                    + clean_rtf(RTF_BULLET)
                    + r" \tab "
                    + clean_rtf(bullet)
                    + r"\par"
                )
        if job.get("tech"):
            line(job["tech"], italic=True)
        else:
            rtf.append(r"\par")

    if data.get("earlier"):
        line(SECTION_EARLIER, bold=True, ul=True)
        for role in data["earlier"]:
            rtf.append(
                r"\pard\sa200\sl276\slmult1\tx360\li360\fi-360 "
                + clean_rtf(RTF_BULLET)
                + r" \tab "
                + clean_rtf(role)
                + r"\par"
            )

    line(SECTION_CERTS_EDU, bold=True, ul=True)
    for item in data["certs"] + data["education"]:
        bold = STRING_IN_PROGRESS in item or STRING_UNIVERSITY in item
        txt = clean_rtf(item)
        if bold:
            txt = r"\b " + txt + r"\b0"
        rtf.append(
            r"\pard\sa200\sl276\slmult1\tx360\li360\fi-360 "
            + clean_rtf(RTF_BULLET)
            + r" \tab "
            + txt
            + r"\par"
        )

    rtf.append("}")
    with open(filename, "w", encoding="utf-8") as f:
        f.write("\n".join(rtf))


# --- MAIN EXECUTION ---
if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Generate Resume from YAML in multiple formats."
    )
    parser.add_argument(
        "--format",
        choices=["docx", "pdf", "md", "rtf", "all"],
        default="docx",
        help="Output format(s)",
    )
    parser.add_argument("--name", default="resume", help="Output filename base")
    parser.add_argument(
        "--source",
        default="../data/example.yml",
        help="Source YAML file (default: ../data/example.yml)",
    )
    parser.add_argument(
        "--reveal-pii", action="store_true", help="Reveal PII (Phone/City) in output"
    )

    # 1. CHECK ARGS: If none provided, print help and exit
    if len(sys.argv) == 1:
        parser.print_help()
        print("\n[-] Note: No arguments provided. Usage examples:")
        print("    python build_resume.py --format all")
        print("    python build_resume.py --format pdf --reveal-pii")
        sys.exit(0)

    args = parser.parse_args()

    # 2. CHECK ENV: Look for .env.local or .env (optional - system env vars also work)
    env_local = ".env.local"
    env_default = ".env"

    if os.path.exists(env_local):
        load_dotenv(env_local, override=True)
        print(f"[+] Loaded configuration from {env_local}")
    elif os.path.exists(env_default):
        load_dotenv(env_default)
        print(f"[+] Loaded configuration from {env_default}")
    else:
        print(f"[*] No .env file found - using system environment variables")
    
    # Debug: Show which PII env vars are available (without revealing values)
    pii_vars = ["RESUME_NAME", "RESUME_CITY_STATE", "RESUME_PHONE", "RESUME_EMAIL", "RESUME_LINKEDIN"]
    found_vars = [v for v in pii_vars if os.getenv(v)]
    missing_vars = [v for v in pii_vars if not os.getenv(v)]
    
    if found_vars:
        print(f"[+] Found environment variables: {', '.join(found_vars)}")
    if missing_vars:
        print(f"[-] Missing environment variables: {', '.join(missing_vars)}")

    # 3. LOAD YAML
    if not os.path.exists(args.source):
        print(f"[!] Error: Source file '{args.source}' not found.", file=sys.stderr)
        sys.exit(1)

    try:
        with open(args.source, "r", encoding="utf-8") as f:
            raw_data = yaml.safe_load(f)
    except yaml.YAMLError as e:
        print(f"[!] Error parsing YAML: {e}", file=sys.stderr)
        sys.exit(1)

    if "resume" not in raw_data:
        print(f"[!] Error: YAML must contain a 'resume' key.", file=sys.stderr)
        sys.exit(1)

    data = raw_data["resume"]
    print(f"[+] Loaded resume data from '{args.source}'")

    # 4. INJECT HEADER (Combine Env + Args)
    data["header"] = {
        "name": os.getenv("RESUME_NAME", "Your Name"),
        "contact": get_contact_string(args.reveal_pii),
    }

    # 5. RENDER
    try:
        if args.format in ["docx", "all"]:
            render_docx(data, f"{args.name}.docx")
        if args.format in ["md", "all"]:
            render_md(data, f"{args.name}.md")
        if args.format in ["pdf", "all"]:
            render_pdf(data, f"{args.name}.pdf")
        if args.format in ["rtf", "all"]:
            render_rtf(data, f"{args.name}.rtf")
        print("[+] Done.")
    except Exception as e:
        print(f"[!] Unexpected error during generation: {e}", file=sys.stderr)
        sys.exit(1)
